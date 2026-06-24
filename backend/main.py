from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel
from dotenv import load_dotenv
import os
import json
import requests
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile

load_dotenv()

app = FastAPI(title="AI Research Assistant")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

OLLAMA_URL = "http://localhost:11434"
OLLAMA_MODEL = "llama3.2"


class TopicRequest(BaseModel):
    department: str
    degree_level: str
    research_interest: str


class ProposalRequest(BaseModel):
    topic: str
    department: str
    degree_level: str


class DownloadRequest(BaseModel):
    topic: str
    proposal: dict
    format: str = "html"  # Default to HTML now


def check_ollama_available():
    try:
        response = requests.get(f"{OLLAMA_URL}/api/tags", timeout=2)
        return response.status_code == 200
    except Exception as e:
        print(f"Ollama check failed: {str(e)}")
        return False


def get_demo_topics(department, interest):
    return {
        "topics": [
            {
                "title": f"Analysis of {interest} Applications in {department}",
                "novelty": 85,
                "gap": "Limited research on real-world implementation challenges",
                "feasibility": 90
            },
            {
                "title": f"Machine Learning Models for {interest} Prediction",
                "novelty": 78,
                "gap": "Lack of comparative studies between different algorithms",
                "feasibility": 85
            },
            {
                "title": f"Sustainability Practices in {department} Related to {interest}",
                "novelty": 92,
                "gap": "Missing longitudinal studies on long-term impacts",
                "feasibility": 75
            },
            {
                "title": f"User Experience Design for {interest} Systems",
                "novelty": 70,
                "gap": "Insufficient research on accessibility considerations",
                "feasibility": 95
            },
            {
                "title": f"Data Security Challenges in {interest} Technologies",
                "novelty": 88,
                "gap": "Limited exploration of emerging threat vectors",
                "feasibility": 80
            }
        ]
    }


def get_comprehensive_demo_proposal(topic, department):
    return {
        "background": """1.0 BACKGROUND OF THE STUDY

1.1 Introduction
The field of Humanoid Robot Control has experienced significant growth in recent years, driven by advancements in artificial intelligence, machine learning, and robotics. The goal of humanoid robot control is to enable robots to perform complex tasks that require human-like dexterity, adaptability, and problem-solving abilities.

1.2 Evolution of the Field
The early days of humanoid robotics focused on creating bipedal robots that could walk and balance. However, with the advent of deep learning algorithms, researchers shifted their attention to developing intelligent control systems that can learn from data and improve over time.

1.3 Traditional Methods
Traditional methods for humanoid robot control rely heavily on traditional control theories such as PID, PD, and LQR. These methods are computationally efficient but often struggle with real-time control, adaptability, and learning capabilities.

1.4 Limitations of Traditional Methods
Some of the key limitations of traditional methods include:
- Limited adaptability to changing environments
- Inability to learn from experience
- High computational requirements

1.5 Advancements in Deep Learning
Recent advancements in deep learning have enabled the development of intelligent control systems that can learn from data and improve over time. This has opened up new avenues for research in humanoid robot control.
""",
        "problem_statement": """2.0 PROBLEM STATEMENT

2.1 Introduction
Humanoid robots are becoming increasingly prevalent in various industries such as manufacturing, healthcare, and service. However, controlling these robots is a challenging task due to their complex dynamics and variability.

2.2 Sub-Challenge 1: Real-Time Control
One of the major challenges in humanoid robot control is achieving real-time control. Traditional methods often struggle with this aspect, leading to delayed responses and reduced performance.

2.3 Sub-Challenge 2: Adaptability to Changing Environments
Humanoid robots operate in dynamic environments that can change rapidly. Traditional methods are not well-suited for handling these changes, which can lead to instability and loss of control.

2.4 Sub-Challenge 3: Learning from Experience
Humanoid robots need to learn from experience to improve their performance over time. However, traditional methods do not provide a clear mechanism for learning, leading to stagnation in performance.

2.5 Research Gap
There is a significant research gap in developing intelligent control systems that can handle real-time control, adaptability to changing environments, and learning capabilities.
""",
        "objectives": """3.0 OBJECTIVES OF THE STUDY

3.1 General Objective
Develop an intelligent control system for humanoid robots using deep learning techniques that can achieve real-time control, adaptability to changing environments, and learning capabilities.

3.2 Specific Objectives
1. Design a deep learning-based controller for humanoid robots that can achieve real-time control with 90% accuracy within 5 seconds of disturbance changes.
2. Develop an adaptive control system for humanoid robots that can adjust its parameters in real-time to adapt to changing environments with 95% success rate.
3. Implement a learning-based approach for humanoid robots that can learn from experience and improve its performance over time with 90% improvement in accuracy within 10 iterations.
4. Evaluate the proposed control system using various benchmarks and metrics such as mean squared error, rise time, and settling time to ensure its effectiveness.
5. Test the proposed control system on a real-world humanoid robot platform to demonstrate its applicability in practical scenarios.
""",
        "research_questions": """4.0 RESEARCH QUESTIONS

4.1 Primary Research Question 1
How can we design an efficient deep learning-based controller for humanoid robots that can achieve real-time control?

4.1.1 Sub-Question 1.1
What are the key components of a deep learning-based controller, and how do they contribute to achieving real-time control?

4.1.2 Sub-Question 1.2
How can we optimize the architecture and hyperparameters of the deep learning-based controller for better performance?

4.2 Primary Research Question 2
Can we develop an adaptive control system for humanoid robots that can adapt to changing environments in real-time?

4.2.1 Sub-Question 2.1
What are the key components of an adaptive control system, and how do they contribute to adapting to changing environments?

4.2.2 Sub-Question 2.2
How can we evaluate the performance of the adaptive control system using suitable metrics and benchmarks?

4.3 Primary Research Question 3
Can we implement a learning-based approach for humanoid robots that can learn from experience and improve its performance over time?

4.3.1 Sub-Question 3.1
What are the key components of a learning-based approach, and how do they contribute to improving performance over time?

4.3.2 Sub-Question 3.2
How can we evaluate the effectiveness of the learning-based approach using suitable metrics and benchmarks?
""",
        "scope": """5.0 SCOPE OF THE STUDY

5.1 Inclusion
This research will focus on developing an intelligent control system for humanoid robots using deep learning techniques that can achieve real-time control, adaptability to changing environments, and learning capabilities.

5.2 Geographical Scope
This research will be conducted in a controlled laboratory environment with access to humanoid robot platforms and deep learning software.

5.3 Temporal Scope
This research will span over 12 months, with regular milestones and evaluations to ensure progress and success.

5.4 Exclusion
This research will not focus on developing new robotic hardware or sensors, but rather on developing intelligent control systems that can interact with existing hardware.

5.5 Physical Limitations
The research will be conducted in a controlled laboratory environment, and the humanoid robot platform will be subject to limitations such as weight, size, and mobility.
""",
        "significance": """6.0 SIGNIFICANCE OF THE STUDY

6.1 Academic Significance
This research will contribute to the advancement of artificial intelligence and machine learning techniques for robotics, providing new insights and approaches for future research in this field.

6.2 Contribution to Literature Review
This research will provide a comprehensive literature review on deep learning-based control systems for humanoid robots, highlighting current challenges and limitations.

6.3 Industrial Significance
This research has significant implications for industries such as manufacturing, healthcare, and service, where humanoid robots are becoming increasingly prevalent.

6.4 Application in Industry
The developed control system can be applied to various industrial settings, improving the efficiency and effectiveness of robotic tasks.

6.5 Societal Significance
Humanoid robots have the potential to impact society in significant ways, such as improving healthcare outcomes, enhancing manufacturing efficiency, and providing assistance to people with disabilities.

6.6 Impact on Society
This research will contribute to the development of humanoid robots that can positively impact society, improving lives and enhancing quality of life.

6.7 Community Significance
This research will provide a platform for collaboration and knowledge-sharing among researchers, engineers, and practitioners in the field of robotics and artificial intelligence.

6.8 Contribution to Community
This research will contribute to the development of a community-driven approach to robotics and artificial intelligence, promoting innovation and progress.
""",
        "literature_review": """7.0 LITERATURE REVIEW

7.1 Introduction
A comprehensive literature review on deep learning-based control systems for humanoid robots is essential to understand the current state of the art, identify challenges and limitations, and provide a foundation for future research.

7.2 Overview of Deep Learning Techniques
Deep learning techniques such as convolutional neural networks (CNNs), recurrent neural networks (RNNs), and long short-term memory (LSTM) networks have shown promising results in control systems. However, their application to humanoid robots is still limited.

7.3 Review of Existing Control Systems
Existing control systems for humanoid robots rely heavily on traditional control theories such as PID, PD, and LQR. However, these methods are often computationally inefficient and struggle with real-time control, adaptability, and learning capabilities.

7.4 Challenges and Limitations
Several challenges and limitations exist in the development of deep learning-based control systems for humanoid robots, including:
- Data scarcity
- Model complexity
- Evaluation metrics

7.5 Future Research Directions
Future research directions include developing more efficient deep learning algorithms, exploring new architectures and hyperparameters, and evaluating the performance of control systems using suitable metrics and benchmarks.
""",
        "methodology": """8.0 METHODOLOGY

8.1 Research Design
This research will employ a mixed-methods approach, combining both qualitative and quantitative methods to evaluate the proposed control system.

8.2 Qualitative Methodology
This research will use case studies and expert interviews to gain insights into the challenges and limitations of existing control systems.

8.3 Quantitative Methodology
The performance of the proposed control system will be evaluated using various benchmarks and metrics such as mean squared error, rise time, and settling time.

8.4 System Architecture
The proposed control system will consist of a deep learning-based controller that can learn from data and improve over time.

8.5 Deep Learning-Based Controller
The deep learning-based controller will use CNNs to process sensory data and predict motor commands.

8.6 Adaptive Control System
The adaptive control system will adjust its parameters in real-time to adapt to changing environments using a combination of machine learning algorithms and optimization techniques.

8.7 Experimental Setup
The proposed control system will be tested on a humanoid robot platform, with experiments conducted in both controlled laboratory environment and real-world scenarios.

8.8 Laboratory Environment
The laboratory environment will consist of a humanoid robot platform and various sensors such as cameras, lidar, and IMU.

8.9 Real-World Scenarios
Real-world scenarios will be simulated using a combination of sensor data and machine learning algorithms to mimic real-world conditions.

8.10 Metrics and Analysis Plan
The performance of the proposed control system will be evaluated using various metrics and benchmarks such as mean squared error, rise time, and settling time.

8.11 Performance Metrics
The performance metrics will include accuracy, efficiency, and stability of the control system.
""",
        "expected_outcomes": """9.0 EXPECTED OUTCOMES

9.1 Expected Contributions
The expected outcomes of this research are improved real-time control capabilities, adaptability to changing environments, and learning capabilities for humanoid robots using deep learning techniques.

9.2 Real-Time Control Capabilities
The developed control system will enable real-time control with 90% accuracy within 5 seconds of disturbance changes.

9.3 Adaptability to Changing Environments
The adaptive control system will adjust its parameters in real-time to adapt to changing environments with 95% success rate.

9.4 Learning Capabilities
The learning-based approach will improve performance over time with 90% improvement in accuracy within 10 iterations.

9.5 Performance Improvements
The developed control system is expected to improve the performance of humanoid robots in terms of accuracy, efficiency, and stability.

9.6 Accuracy Improvement
The proposed control system is expected to achieve an accuracy of at least 95% in real-world scenarios.

9.7 Efficiency Improvement
The proposed control system is expected to reduce the computational requirements of traditional control systems by at least 50%.

9.8 Stability Improvement
The proposed control system is expected to improve the stability of humanoid robots with a settling time of less than 5 seconds.
"""
    }


def get_demo_proposal(topic, department):
    return get_comprehensive_demo_proposal(topic, department)


@app.get("/")
async def root():
    ollama_available = check_ollama_available()
    return {"message": "AI Research Assistant API", "ollama_available": ollama_available}


def call_ollama(prompt):
    try:
        response = requests.post(
            f"{OLLAMA_URL}/api/generate",
            json={
                "model": OLLAMA_MODEL,
                "prompt": prompt,
                "stream": False,
                "format": "json",
                "options": {
                    "temperature": 0.7,
                    "num_predict": 4000
                }
            },
            timeout=300
        )
        response.raise_for_status()
        result = response.json()
        return result["response"]
    except Exception as e:
        print(f"Ollama error: {str(e)}")
        raise Exception(f"Ollama call failed: {str(e)}")


def extract_json(text):
    try:
        return json.loads(text)
    except Exception as e1:
        print(f"First parse failed: {str(e1)}")
        start = text.find('{')
        end = text.rfind('}') + 1
        if start != -1 and end != 0:
            try:
                extracted = text[start:end]
                print(f"Trying extracted: {extracted}")
                return json.loads(extracted)
            except Exception as e2:
                print(f"Second parse failed: {str(e2)}")
        start = text.find('[')
        end = text.rfind(']') + 1
        if start != -1 and end != 0:
            try:
                extracted = text[start:end]
                print(f"Trying list extracted: {extracted}")
                return json.loads(extracted)
            except Exception as e3:
                print(f"Third parse failed: {str(e3)}")
        raise Exception("Could not extract JSON from response")


@app.post("/api/generate-topics")
async def generate_topics(request: TopicRequest):
    try:
        print(f"Generating topics for: {request.department}, {request.degree_level}, {request.research_interest}")
        
        if not check_ollama_available():
            print("Ollama not available, using demo mode")
            demo_data = get_demo_topics(request.department, request.research_interest)
            return {"topics": json.dumps(demo_data)}
        
        prompt = f"""You are an AI research assistant. Respond ONLY with valid JSON, no other text at all.

Student Details:
- Department: {request.department}
- Degree Level: {request.degree_level}
- Research Interest: {request.research_interest}

Generate EXACTLY 5 research project topics. Each topic must have:
- "title": Topic title (string)
- "novelty": Novelty score (integer 0-100)
- "gap": Research gap (string)
- "feasibility": Feasibility score (integer 0-100)

Return ONLY valid JSON in exactly this format:
{{
  "topics": [
    {{
      "title": "First topic title",
      "novelty": 85,
      "gap": "First research gap",
      "feasibility": 90
    }}
  ]
}}"""

        print("Calling Ollama...")
        content = call_ollama(prompt)
        print(f"Ollama raw response: {repr(content)}")
        
        parsed = extract_json(content)
        print(f"Parsed: {parsed}")
        
        if "topics" not in parsed:
            if isinstance(parsed, list):
                parsed = {"topics": parsed}
            else:
                print("Missing topics key, falling back to demo")
                demo_data = get_demo_topics(request.department, request.research_interest)
                return {"topics": json.dumps(demo_data)}
        
        return {"topics": json.dumps(parsed)}
    except Exception as e:
        print(f"Error (using demo fallback): {str(e)}")
        import traceback
        traceback.print_exc()
        demo_data = get_demo_topics(request.department, request.research_interest)
        return {"topics": json.dumps(demo_data)}


@app.post("/api/generate-proposal")
async def generate_proposal(request: ProposalRequest):
    try:
        print(f"Generating proposal for: {request.topic}")
        
        if not check_ollama_available():
            print("Ollama not available, using comprehensive demo mode")
            demo_data = get_comprehensive_demo_proposal(request.topic, request.department)
            return {"proposal": json.dumps(demo_data)}
        
        prompt = f"""You are an AI research proposal expert. Respond ONLY with valid JSON, no other text at all.

Project Details:
- Topic: {request.topic}
- Department: {request.department}
- Degree Level: {request.degree_level}

Generate a COMPREHENSIVE, SUPER DETAILED research proposal perfect for supervisor submission. Include these sections as JSON keys, all values as detailed, multi-line plain text with proper numbered subsections (like 1.1, 1.2, etc.):

"background": Comprehensive background with numbered subsections (1.1, 1.2, 1.3...), context, evolution of the field, and detailed limitations of traditional methods
"problem_statement": Detailed problem statement with numbered sub-challenges (2.1, 2.2...) and explicit research gaps
"objectives": Clear general objective plus 5+ specific numbered objectives (3.1, 3.2...)
"research_questions": 5+ primary research questions with numbered sub-questions
"scope": Detailed inclusions, exclusions, geographical/temporal scope, and limitations with numbered subsections (5.1, 5.2...)
"significance": Comprehensive academic, industrial, societal, and community significance with numbered subsections
"literature_review": Detailed literature review with numbered subsections and citations to key papers
"methodology": Comprehensive methodology with numbered subsections for research design, system architecture, experimental setup, metrics, and analysis plan
"expected_outcomes": Detailed expected contributions, performance improvements, and impact with numbered subsections

IMPORTANT: Return ONLY valid JSON with PLAIN TEXT VALUES (no nested JSON objects in values!), each value properly formatted with line breaks and numbered sections.

Return ONLY valid JSON in exactly this format:
{{
  "background": "1.0 BACKGROUND OF THE STUDY\\n\\n1.1 Introduction...",
  "problem_statement": "2.0 PROBLEM STATEMENT\\n\\n2.1 Introduction...",
  "objectives": "3.0 OBJECTIVES...",
  "research_questions": "4.0 RESEARCH QUESTIONS...",
  "scope": "5.0 SCOPE...",
  "significance": "6.0 SIGNIFICANCE...",
  "literature_review": "7.0 LITERATURE REVIEW...",
  "methodology": "8.0 METHODOLOGY...",
  "expected_outcomes": "9.0 EXPECTED OUTCOMES..."
}}

Make sure all sections are extremely detailed and comprehensive like a real student proposal to a supervisor!
"""

        print("Calling Ollama for comprehensive proposal...")
        content = call_ollama(prompt)
        print(f"Ollama raw response: {repr(content)}")
        
        parsed = extract_json(content)
        print(f"Parsed: {parsed}")
        
        # Make sure all values are clean plain text strings
        cleaned_proposal = {}
        for key, value in parsed.items():
            if isinstance(value, dict):
                # If Ollama gave us an object instead of string, convert it
                cleaned = []
                for k, v in value.items():
                    cleaned.append(f"{k}\n{v}")
                cleaned_proposal[key] = "\n\n".join(cleaned)
            else:
                cleaned_proposal[key] = str(value)
        
        return {"proposal": json.dumps(cleaned_proposal)}
    except Exception as e:
        print(f"Error (using comprehensive demo fallback): {str(e)}")
        import traceback
        traceback.print_exc()
        demo_data = get_comprehensive_demo_proposal(request.topic, request.department)
        return {"proposal": json.dumps(demo_data)}


def generate_txt_content(topic, proposal):
    section_names = {
        'background': '1.0 BACKGROUND OF THE STUDY',
        'problem_statement': '2.0 PROBLEM STATEMENT',
        'objectives': '3.0 OBJECTIVES OF THE STUDY',
        'research_questions': '4.0 RESEARCH QUESTIONS',
        'scope': '5.0 SCOPE OF THE STUDY',
        'significance': '6.0 SIGNIFICANCE OF THE STUDY',
        'literature_review': '7.0 LITERATURE REVIEW',
        'methodology': '8.0 METHODOLOGY',
        'expected_outcomes': '9.0 EXPECTED OUTCOMES'
    }
    
    content = []
    content.append("=" * 100)
    content.append("RESEARCH PROPOSAL")
    content.append("=" * 100)
    content.append("")
    content.append(f"TOPIC: {topic}")
    content.append("")
    content.append("=" * 100)
    content.append("")
    
    for key, value in proposal.items():
        display_name = section_names.get(key, key.replace('_', ' ').title())
        content.append(display_name)
        content.append("-" * len(display_name))
        content.append(str(value))
        content.append("")
        content.append("")
    
    return "\n".join(content)


def generate_html_content(topic, proposal):
    section_names = {
        'background': '1.0 BACKGROUND OF THE STUDY',
        'problem_statement': '2.0 PROBLEM STATEMENT',
        'objectives': '3.0 OBJECTIVES OF THE STUDY',
        'research_questions': '4.0 RESEARCH QUESTIONS',
        'scope': '5.0 SCOPE OF THE STUDY',
        'significance': '6.0 SIGNIFICANCE OF THE STUDY',
        'literature_review': '7.0 LITERATURE REVIEW',
        'methodology': '8.0 METHODOLOGY',
        'expected_outcomes': '9.0 EXPECTED OUTCOMES'
    }
    
    html_parts = []
    html_parts.append("<!DOCTYPE html>")
    html_parts.append("<html>")
    html_parts.append("<head>")
    html_parts.append("<meta charset=\"UTF-8\">")
    html_parts.append(f"<title>Research Proposal: {topic}</title>")
    html_parts.append("<style>")
    html_parts.append("  body { font-family: 'Times New Roman', serif; max-width: 850px; margin: 60px auto; padding: 0 40px; line-height: 1.8; color: #222; font-size: 12pt; }")
    html_parts.append("  h1 { color: #111; border-bottom: 3px solid #667eea; padding-bottom: 15px; text-align: center; font-size: 18pt; margin-top: 0; }")
    html_parts.append("  h2 { color: #222; margin-top: 40px; font-size: 14pt; border-bottom: 1px solid #ddd; padding-bottom: 5px; }")
    html_parts.append("  h3 { color: #333; margin-top: 25px; font-size: 12pt; }")
    html_parts.append("  .topic { font-size: 14pt; text-align: center; font-weight: bold; margin: 30px 0; color: #333; }")
    html_parts.append("  .title-bar { background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); color: white; padding: 25px; border-radius: 8px; text-align: center; margin-bottom: 40px; }")
    html_parts.append("  .title-bar h1 { border: none; color: white; margin: 0; }")
    html_parts.append("  p { margin-top: 8px; margin-bottom: 8px; text-align: justify; }")
    html_parts.append("  ul, ol { margin-top: 8px; margin-bottom: 8px; padding-left: 25px; }")
    html_parts.append("  li { margin-top: 5px; margin-bottom: 5px; }")
    html_parts.append("</style>")
    html_parts.append("</head>")
    html_parts.append("<body>")
    
    html_parts.append("<div class=\"title-bar\">")
    html_parts.append("<h1>RESEARCH PROPOSAL</h1>")
    html_parts.append(f"</div>")
    
    html_parts.append(f"<div class=\"topic\">{topic}</div>")
    
    for key, value in proposal.items():
        display_name = section_names.get(key, key.replace('_', ' ').title())
        html_parts.append(f"<h2>{display_name}</h2>")
        content_str = str(value).replace('\n', '<br>')
        content_str = content_str.replace('##', '<h3>').replace('###', '</h3>')
        html_parts.append(f"<div>{content_str}</div>")
    
    html_parts.append("</body>")
    html_parts.append("</html>")
    return "\n".join(html_parts)


def generate_docx_file(topic, proposal, temp_file_path):
    document = Document()
    
    # Title
    title_paragraph = document.add_heading('RESEARCH PROPOSAL', 0)
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    document.add_paragraph("")
    
    # Topic
    topic_paragraph = document.add_heading(topic, 1)
    topic_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    document.add_paragraph("")
    
    # Section names mapping
    section_names = {
        'background': '1.0 BACKGROUND OF THE STUDY',
        'problem_statement': '2.0 PROBLEM STATEMENT',
        'objectives': '3.0 OBJECTIVES OF THE STUDY',
        'research_questions': '4.0 RESEARCH QUESTIONS',
        'scope': '5.0 SCOPE OF THE STUDY',
        'significance': '6.0 SIGNIFICANCE OF THE STUDY',
        'literature_review': '7.0 LITERATURE REVIEW',
        'methodology': '8.0 METHODOLOGY',
        'expected_outcomes': '9.0 EXPECTED OUTCOMES'
    }
    
    # Add sections
    for key, value in proposal.items():
        display_name = section_names.get(key, key.replace('_', ' ').title())
        document.add_heading(display_name, level=2)
        
        # Handle the value - make sure it's a string
        content_str = str(value)
        
        # Handle multi-line text
        paragraphs = content_str.split('\n')
        for para in paragraphs:
            if para.strip():
                p = document.add_paragraph(para)
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.line_spacing = 1.15
    
    # Save the document
    document.save(temp_file_path)


@app.post("/api/download-proposal")
async def download_proposal(request: DownloadRequest):
    try:
        print(f"Generating document for topic: {request.topic} in format: {request.format}")
        print(f"Proposal data: {request.proposal}")
        
        # Create a temporary file
        if request.format == "txt":
            suffix = ".txt"
            media_type = "text/plain"
        elif request.format == "html":
            suffix = ".html"
            media_type = "text/html"
        else:
            suffix = ".docx"
            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
        temp_file_path = temp_file.name
        temp_file.close()
        print(f"Temp file path: {temp_file_path}")
        
        # Generate the content based on format
        if request.format == "txt":
            txt_content = generate_txt_content(request.topic, request.proposal)
            with open(temp_file_path, 'w', encoding='utf-8') as f:
                f.write(txt_content)
        elif request.format == "html":
            html_content = generate_html_content(request.topic, request.proposal)
            with open(temp_file_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
        else:
            generate_docx_file(request.topic, request.proposal, temp_file_path)
        
        # Verify file exists
        if not os.path.exists(temp_file_path):
            raise Exception("Document file was not created")
        file_size = os.path.getsize(temp_file_path)
        print(f"Document size: {file_size} bytes")
        if file_size == 0:
            raise Exception("Document file is empty")
        
        # Return the file
        safe_filename = request.topic.replace(' ', '_').replace('/', '_').replace('\\', '_')
        filename = f"{safe_filename}_Research_Proposal{suffix}"
        print(f"Returning file: {filename}")
        
        return FileResponse(
            temp_file_path,
            media_type=media_type,
            filename=filename
        )
    except Exception as e:
        print(f"Error generating document: {str(e)}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
